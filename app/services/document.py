"""
Генерация документов .docx из шаблонов.
Логика полностью перенесена из оригинального main.py без изменений.
"""
import copy
import io
import logging
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# 1x1 прозрачный PNG для скрытия блока подписи руководителя
_TRANSPARENT_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


def set_approver_row_visibility(doc: Document, visible: bool) -> None:
    """
    Скрывает или показывает строку таблицы с блоком руководителя (SIGN_APPROVER).
    visible=False — красим текст в белый, картинку делаем прозрачной.
    visible=True  — возвращаем чёрный цвет, картинку заменит replace_image.
    """
    from lxml import etree
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    color = "FFFFFF" if not visible else "000000"

    for table in doc.tables:
        for row in table.rows:
            row_xml = etree.tostring(row._element, encoding="unicode")
            if "SIGN_APPROVER" not in row_xml:
                continue

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is None:
                            from lxml import etree as _etree
                            rpr = _etree.SubElement(run._element, qn("w:rPr"))
                            run._element.insert(0, rpr)
                        color_el = rpr.find(qn("w:color"))
                        if color_el is None:
                            from lxml import etree as _etree
                            color_el = _etree.SubElement(rpr, qn("w:color"))
                        color_el.set(qn("w:val"), color)

            for drawing in row._element.findall(".//" + qn("wp:anchor")):
                docPr = drawing.find(".//" + qn("wp:docPr"))
                if docPr is None or docPr.get("descr") != "SIGN_APPROVER":
                    continue
                blip = drawing.find(".//" + qn("a:blip"))
                if blip is None:
                    continue
                if not visible:
                    transp_part = Part(
                        partname=PackURI("/word/media/sign_approver_hidden.png"),
                        content_type="image/png",
                        blob=_TRANSPARENT_PNG,
                        package=doc.part.package,
                    )
                    new_rId = doc.part.relate_to(
                        transp_part,
                        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
                    )
                    blip.set(qn("r:embed"), new_rId)
                    logger.info("SIGN_APPROVER скрыт (прозрачный PNG)")
            return


def _replace_paragraph_text(paragraph, data: dict) -> None:
    """
    Заменяет переменные вида ${KEY} в одном параграфе.
    Собирает полный текст из всех run-ов чтобы обойти разбивку Word.
    Если значение содержит \\n — разбивает параграф на несколько.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    original_text = full_text

    for key, value in data.items():
        placeholder = f"${{{key}}}"
        if placeholder in full_text:
            str_value = str(value)
            if str_value.startswith("\n") and f"{placeholder}." in full_text:
                full_text = full_text.replace(f"{placeholder}.", str_value)
            else:
                full_text = full_text.replace(placeholder, str_value)

    if full_text == original_text:
        return

    if "\n" not in full_text:
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""
        return

    lines = full_text.split("\n")
    if paragraph.runs:
        paragraph.runs[0].text = lines[0]
        for run in paragraph.runs[1:]:
            run.text = ""

    parent = paragraph._element.getparent()
    insert_idx = list(parent).index(paragraph._element)
    for i, line in enumerate(lines[1:], 1):
        new_p = copy.deepcopy(paragraph._element)
        runs_in_new_p = new_p.findall(".//" + qn("w:r"))
        for r in runs_in_new_p[1:]:
            r.getparent().remove(r)
        if runs_in_new_p:
            t = runs_in_new_p[0].find(qn("w:t"))
            if t is not None:
                t.text = line
                if line and (line[0] == " " or line[-1] == " "):
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        parent.insert(insert_idx + i, new_p)


def replace_text(doc: Document, data: dict) -> None:
    """Заменяет текстовые переменные во всём документе — в параграфах и таблицах."""
    for paragraph in list(doc.paragraphs):
        _replace_paragraph_text(paragraph, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    _replace_paragraph_text(paragraph, data)


def replace_image(doc: Document, placeholder_desc: str, image_bytes: bytes) -> bool:
    """
    Заменяет картинку-заглушку на реальное изображение подписи.
    Поиск по полю Alt Text (атрибут descr в wp:docPr).
    """
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    def _process_paragraphs(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                drawings = (
                    run._element.findall(".//" + qn("wp:inline"))
                    + run._element.findall(".//" + qn("wp:anchor"))
                )
                for drawing in drawings:
                    docPr = drawing.find(".//" + qn("wp:docPr"))
                    if docPr is None or docPr.get("descr", "") != placeholder_desc:
                        continue
                    blip = drawing.find(".//" + qn("a:blip"))
                    if blip is None:
                        continue
                    image_part = Part(
                        partname=PackURI(f"/word/media/sign_{placeholder_desc}.png"),
                        content_type="image/png",
                        blob=image_bytes,
                        package=doc.part.package,
                    )
                    new_rId = doc.part.relate_to(
                        image_part,
                        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
                    )
                    blip.set(qn("r:embed"), new_rId)
                    logger.info(f"Заменена картинка: {placeholder_desc} -> {new_rId}")
                    return True
        return False

    if _process_paragraphs(doc.paragraphs):
        return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if _process_paragraphs(cell.paragraphs):
                    return True
    return False


def parse_variables(variables_list) -> dict:
    """Парсит переменные из формата KEY|VALUE в словарь."""
    result = {}
    if isinstance(variables_list, str):
        variables_list = [variables_list]
    for item in variables_list or []:
        if "|" in item:
            key, _, value = item.partition("|")
            result[key.strip()] = value.strip()
    return result


def parse_signatures(signatures_list) -> list:
    """Парсит подписи из формата PLACEHOLDER|FILE_ID в список словарей."""
    result = []
    if isinstance(signatures_list, str):
        signatures_list = [signatures_list]
    for item in signatures_list or []:
        if "|" in item:
            placeholder, _, file_id = item.partition("|")
            result.append({"placeholder": placeholder.strip(), "signature_id": file_id.strip()})
    return result


def format_request_goal(raw_goal: str) -> str:
    """
    Форматирует REQUEST_GOAL из сырых данных "Назначение|Сумма\\nНазначение2|Сумма2".
    Одна позиция → просто название. Несколько → нумерованный список.
    """
    lines = [l.strip() for l in raw_goal.split("\n") if l.strip()]
    if len(lines) == 1:
        return lines[0].split("|")[0].strip()

    formatted = []
    for i, line in enumerate(lines, 1):
        parts = line.split("|")
        name = parts[0].strip()
        amount = parts[1].strip() if len(parts) > 1 else ""
        if amount:
            formatted.append(f"{i}. {name} — {int(float(amount)):,} руб.".replace(",", " "))
        else:
            formatted.append(f"{i}. {name}")
    return "\n" + "\n".join(formatted)


async def generate_document(
    bitrix,
    template_id: str,
    folder_id: str,
    filename: str,
    data: dict,
    signatures: list,
    source_file_id: str = "",
) -> str:
    """
    Генерирует документ из шаблона и загружает на Диск Bitrix24.
    Если source_file_id передан — берёт существующий файл (для добавления подписи руководителя).
    Возвращает ID загруженного файла.
    """
    file_id_to_use = source_file_id if source_file_id else template_id
    file_bytes = await bitrix.download_file(file_id_to_use)
    doc = Document(io.BytesIO(file_bytes))

    if not source_file_id:
        # Первая генерация — подставляем переменные и скрываем блок руководителя
        replace_text(doc, data)
        set_approver_row_visibility(doc, visible=False)
    else:
        # После согласования — показываем блок руководителя
        set_approver_row_visibility(doc, visible=True)

    for sig in signatures:
        sign_bytes = await bitrix.download_file(sig["signature_id"])
        replace_image(doc, sig["placeholder"], sign_bytes)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    name, ext = filename.rsplit(".", 1) if "." in filename else (filename, "docx")
    unique_filename = f"{name}_{timestamp}.{ext}"

    result_id = await bitrix.upload_file(folder_id, unique_filename, output.read())
    logger.info(f"Документ сгенерирован: ID={result_id}, файл={unique_filename}")
    return result_id
