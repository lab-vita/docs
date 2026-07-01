import os
import io
import json
import requests
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
from typing import List, Optional
from docx import Document
from docx.oxml.ns import qn
from datetime import datetime
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Умные бизнес-процессы")

CLIENT_ID = os.getenv("CLIENT_ID")
CLIENT_SECRET = os.getenv("CLIENT_SECRET")
BITRIX_WEBHOOK = os.getenv("BITRIX_WEBHOOK_URL")
APP_URL = os.getenv("APP_URL", "https://docs.lab-vita.ru")
TOKEN_FILE = "/app/tokens.json"

CASH_REQUEST_IBLOCK_ID = os.getenv("CASH_REQUEST_IBLOCK_ID", "94")
CASH_REQUEST_BP_ID = os.getenv("CASH_REQUEST_BP_ID", "592")
FIELD_SUMMA = os.getenv("FIELD_SUMMA", "PROPERTY_516")
FIELD_NAZNACHENIE = os.getenv("FIELD_NAZNACHENIE", "PROPERTY_518")

APP_TITLE = "Умные бизнес-процессы"

# === Реестр процессов ===
# Единая точка регистрации доступных бизнес-процессов.
# Чтобы добавить новый процесс — добавь запись в этот словарь
# и соответствующий form_*.html в той же директории.

PROCESSES = {
    "cash": {
        "title": "Выдача наличных",
        "description": "Заявка на выдачу денежных средств под отчёт",
        "icon": "💵",
        "form_file": "form_cash.html",
    },
}


# === Модели данных ===

class SignatureEntry(BaseModel):
    placeholder: str  # Замещающий текст картинки-заглушки в шаблоне
    signature_id: str  # ID файла подписи на Диске Битрикс24


class Position(BaseModel):
    name: str
    amount: float


class SubmitRequest(BaseModel):
    title: str
    positions: List[Position]
    user_id: str = ""
    auth_token: str = ""
    user_auth_id: str = ""  # Токен пользователя — используется для создания элемента от его имени


# === Работа с токенами ===

def save_tokens(tokens: dict):
    """Сохраняет токены OAuth в файл"""
    with open(TOKEN_FILE, "w") as f:
        json.dump(tokens, f)
    logger.info("Токены сохранены")


def load_tokens() -> dict:
    """Загружает токены OAuth из файла"""
    try:
        with open(TOKEN_FILE, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}


def refresh_access_token() -> str:
    """Обновляет access_token через refresh_token и возвращает новый"""
    tokens = load_tokens()
    if not tokens.get("refresh_token"):
        raise HTTPException(status_code=401, detail="Нет refresh_token — переустановите приложение")

    resp = requests.post("https://oauth.bitrix.info/oauth/token/", params={
        "grant_type": "refresh_token",
        "client_id": CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "refresh_token": tokens["refresh_token"],
    })
    resp.raise_for_status()
    new_tokens = resp.json()
    save_tokens(new_tokens)
    logger.info("Токен обновлён")
    return new_tokens["access_token"]


def get_access_token() -> str:
    """Возвращает актуальный access_token из файла"""
    tokens = load_tokens()
    if not tokens.get("access_token"):
        return refresh_access_token()
    return tokens["access_token"]


def get_client_endpoint() -> str:
    """Возвращает базовый URL REST API Битрикс24"""
    tokens = load_tokens()
    return tokens.get("client_endpoint", "")


# === Интеграция с Диском Битрикс24 ===

def b24_download_file(file_id: str) -> bytes:
    """Скачивает файл с Диска Битрикс24 по ID через токен приложения"""
    access_token = get_access_token()
    client_endpoint = get_client_endpoint()
    url = f"{client_endpoint}disk.file.get.json"
    resp = requests.post(url, params={"auth": access_token}, json={"id": file_id})
    resp.raise_for_status()
    result = resp.json().get("result", {})
    download_url = result.get("DOWNLOAD_URL")
    if not download_url:
        raise HTTPException(status_code=404, detail=f"Файл {file_id} не найден на Диске")
    file_resp = requests.get(download_url)
    file_resp.raise_for_status()
    return file_resp.content


def b24_upload_file(folder_id: str, filename: str, content: bytes) -> str:
    """Загружает файл на Диск Битрикс24 через вебхук и возвращает ID"""
    url = f"{BITRIX_WEBHOOK}disk.folder.uploadfile.json"
    resp = requests.post(url, params={"id": folder_id, "data[NAME]": filename})
    resp.raise_for_status()
    upload_url = resp.json().get("result", {}).get("uploadUrl")
    if not upload_url:
        raise HTTPException(status_code=500, detail="Не получен upload URL")

    files = {"file": (filename, content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    upload_resp = requests.post(upload_url, files=files)
    upload_resp.raise_for_status()
    file_id = upload_resp.json().get("result", {}).get("ID")
    if not file_id:
        raise HTTPException(status_code=500, detail="Не получен ID файла")
    return str(file_id)


# === Обработка документа ===

def set_approver_row_visibility(doc: Document, visible: bool):
    """
    Скрывает или показывает строку таблицы с блоком руководителя (SIGN_APPROVER).
    visible=False — красим текст в белый, картинку делаем прозрачной (1x1 пиксель).
    visible=True  — возвращаем чёрный цвет текста, картинку заменит replace_image.
    """
    from lxml import etree

    # 1x1 прозрачный PNG
    TRANSPARENT_PNG = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
        b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
    )

    color = "FFFFFF" if not visible else "000000"

    for table in doc.tables:
        for row in table.rows:
            # Ищем строку содержащую SIGN_APPROVER
            row_xml = etree.tostring(row._element, encoding="unicode")
            if "SIGN_APPROVER" not in row_xml:
                continue

            # Красим весь текст в строке
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is None:
                            rpr = etree.SubElement(run._element, qn("w:rPr"))
                            run._element.insert(0, rpr)
                        color_el = rpr.find(qn("w:color"))
                        if color_el is None:
                            color_el = etree.SubElement(rpr, qn("w:color"))
                        color_el.set(qn("w:val"), color)

            # Скрываем/показываем картинку SIGN_APPROVER
            for drawing in row._element.findall(".//" + qn("wp:anchor")):
                docPr = drawing.find(".//" + qn("wp:docPr"))
                if docPr is None or docPr.get("descr") != "SIGN_APPROVER":
                    continue
                blip = drawing.find(".//" + qn("a:blip"))
                if blip is None:
                    continue
                if not visible:
                    # Заменяем на прозрачный PNG
                    from docx.opc.part import Part
                    from docx.opc.packuri import PackURI
                    transp_part = Part(
                        partname=PackURI("/word/media/sign_approver_hidden.png"),
                        content_type="image/png",
                        blob=TRANSPARENT_PNG,
                        package=doc.part.package
                    )
                    new_rId = doc.part.relate_to(
                        transp_part,
                        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
                    )
                    blip.set(qn("r:embed"), new_rId)
                    logger.info("SIGN_APPROVER скрыт (прозрачный PNG)")
                # visible=True — картинку заменит replace_image после этой функции
            return

def replace_paragraph_text(paragraph, data: dict):
    """
    Заменяет переменные вида ${KEY} в одном параграфе.
    Собирает полный текст из всех run-ов чтобы обойти разбивку Word.
    Если значение содержит \n — разбивает параграф на несколько.
    """
    import copy

    # Сначала собираем полный текст параграфа и делаем все замены за один проход
    full_text = "".join(run.text for run in paragraph.runs)
    original_text = full_text

    for key, value in data.items():
        placeholder = f"${{{key}}}"
        if placeholder in full_text:
            str_value = str(value)
            # Если значение начинается с \n (список позиций),
            # убираем точку стоящую сразу после плейсхолдера чтобы не было "на ."
            if str_value.startswith("\n") and f"{placeholder}." in full_text:
                full_text = full_text.replace(f"{placeholder}.", str_value)
            else:
                full_text = full_text.replace(placeholder, str_value)

    # Ничего не изменилось — выходим
    if full_text == original_text:
        return

    # Если нет переносов строк — простая замена
    if "\n" not in full_text:
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""
        return

    # Есть переносы — разбиваем на несколько параграфов
    lines = full_text.split("\n")

    # Первую строку вставляем в текущий параграф
    if paragraph.runs:
        paragraph.runs[0].text = lines[0]
        for run in paragraph.runs[1:]:
            run.text = ""

    # Остальные строки вставляем как новые параграфы после текущего
    parent = paragraph._element.getparent()
    insert_idx = list(parent).index(paragraph._element)
    for i, line in enumerate(lines[1:], 1):
        new_p = copy.deepcopy(paragraph._element)
        runs_in_new_p = new_p.findall('.//' + qn('w:r'))
        # Удаляем лишние runs, оставляем только первый
        for r in runs_in_new_p[1:]:
            r.getparent().remove(r)
        if runs_in_new_p:
            t = runs_in_new_p[0].find(qn('w:t'))
            if t is not None:
                t.text = line
                if line and (line[0] == ' ' or line[-1] == ' '):
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        parent.insert(insert_idx + i, new_p)


def replace_text(doc: Document, data: dict):
    """Заменяет текстовые переменные во всём документе — в параграфах и таблицах"""
    for paragraph in list(doc.paragraphs):
        replace_paragraph_text(paragraph, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    replace_paragraph_text(paragraph, data)


def replace_image(doc: Document, placeholder_desc: str, image_bytes: bytes):
    """
    Заменяет картинку-заглушку на реальное изображение подписи.
    Поиск по полю "Описание" замещающего текста (Alt Text).
    Создаёт новый независимый part чтобы разные подписи не шарили один файл.
    """
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    def process_paragraphs(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                drawings = (
                        run._element.findall('.//' + qn('wp:inline')) +
                        run._element.findall('.//' + qn('wp:anchor'))
                )
                for drawing in drawings:
                    docPr = drawing.find('.//' + qn('wp:docPr'))
                    if docPr is None:
                        continue
                    if docPr.get('descr', '') != placeholder_desc:
                        continue
                    blip = drawing.find('.//' + qn('a:blip'))
                    if blip is None:
                        continue

                    image_part = Part(
                        partname=PackURI(f"/word/media/sign_{placeholder_desc}.png"),
                        content_type="image/png",
                        blob=image_bytes,
                        package=doc.part.package
                    )
                    new_rId = doc.part.relate_to(
                        image_part,
                        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
                    )
                    blip.set(qn('r:embed'), new_rId)
                    logger.info(f"Заменена картинка: {placeholder_desc} -> {new_rId}")
                    return True
        return False

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)


def parse_variables(variables_list) -> dict:
    """
    Парсит переменные из формата KEY|VALUE в словарь.
    Пример: ["EMPLOYEE_NAME|Егошин Алексей"] -> {"EMPLOYEE_NAME": "Егошин Алексей"}
    """
    result = {}
    if isinstance(variables_list, str):
        variables_list = [variables_list]
    for item in variables_list or []:
        if "|" in item:
            key, _, value = item.partition("|")
            result[key.strip()] = value.strip()
    return result


def parse_signatures(signatures_list) -> list:
    """
    Парсит подписи из формата PLACEHOLDER|FILE_ID в список словарей.
    Пример: ["SIGN_EMPLOYEE|837426"] -> [{"placeholder": "SIGN_EMPLOYEE", "signature_id": "837426"}]
    """
    result = []
    if isinstance(signatures_list, str):
        signatures_list = [signatures_list]
    for item in signatures_list or []:
        if "|" in item:
            placeholder, _, file_id = item.partition("|")
            result.append({"placeholder": placeholder.strip(), "signature_id": file_id.strip()})
    return result


def generate_document(template_id: str, folder_id: str, filename: str, data: dict, signatures: list, source_file_id: str = "") -> str:
    """
    Генерирует документ из шаблона и загружает на Диск.
    Если source_file_id передан — берёт существующий файл вместо шаблона
    (используется для добавления подписи руководителя после согласования).
    Возвращает ID загруженного файла.
    """
    file_id_to_use = source_file_id if source_file_id else template_id
    file_bytes = b24_download_file(file_id_to_use)
    doc = Document(io.BytesIO(file_bytes))

    if not source_file_id:
        # Первая генерация — подставляем переменные и скрываем блок руководителя
        replace_text(doc, data)
        set_approver_row_visibility(doc, visible=False)
    else:
        # После согласования — показываем блок руководителя
        set_approver_row_visibility(doc, visible=True)

    for sig in signatures:
        sign_bytes = b24_download_file(sig["signature_id"])
        replace_image(doc, sig["placeholder"], sign_bytes)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    name, ext = filename.rsplit(".", 1) if "." in filename else (filename, "docx")
    unique_filename = f"{name}_{timestamp}.{ext}"

    file_id = b24_upload_file(folder_id, unique_filename, output.read())
    logger.info(f"Документ сгенерирован, ID: {file_id}, имя: {unique_filename}")
    return file_id


# === Регистрация активити в Битрикс24 ===

def _bind_menu_to_all_users(access_token: str, client_endpoint: str):
    """
    Добавляет пункт левого меню каждому активному пользователю портала.
    Битрикс24 не добавляет placement.bind автоматически для локальных приложений,
    поэтому делаем это программно при установке/переустановке.
    """
    start = 0
    limit = 50
    total_bound = 0

    while True:
        resp = requests.post(f"{client_endpoint}user.get.json", json={
            "auth": access_token,
            "filter": {"ACTIVE": True},
            "select": ["ID"],
            "start": start,
        })
        data = resp.json()
        users = data.get("result", [])
        if not users:
            break

        for user in users:
            user_id = user.get("ID")
            if not user_id:
                continue
            bind_resp = requests.post(f"{client_endpoint}placement.bind.json", json={
                "auth": access_token,
                "PLACEMENT": "LEFT_MENU",
                "HANDLER": f"{APP_URL}/app",
                "TITLE": APP_TITLE,
                "USER_ID": user_id,
            })
            result = bind_resp.json()
            if result.get("result"):
                total_bound += 1

        total = data.get("total", 0)
        start += limit
        if start >= total:
            break

    logger.info(f"Пункт меню добавлен {total_bound} пользователям")

def register_activity(access_token: str):
    """Регистрирует кастомное активити в дизайнере БП и пункт в левом меню"""
    client_endpoint = get_client_endpoint()

    # Удаляем старое активити если уже существует
    requests.post(f"{client_endpoint}bizproc.activity.delete.json", json={
        "auth": access_token,
        "CODE": "generate_document",
    })

    params = {
        "auth": access_token,
        "CODE": "generate_document",
        "HANDLER": f"{APP_URL}/activity-handler",
        "AUTH_USER_ID": 1,
        "USE_SUBSCRIPTION": "Y",
        "NAME": {"ru": "Генерация документа", "en": "Generate Document"},
        "DESCRIPTION": {
            "ru": "Генерирует документ из шаблона .docx с подстановкой переменных и подписей",
            "en": "Generates a document from a .docx template"
        },
        "PROPERTIES": {
            "template_id": {
                "Name": {"ru": "ID шаблона на Диске", "en": "Template ID"},
                "Type": "string",
                "Required": "Y",
                "Multiple": "N",
            },
            "folder_id": {
                "Name": {"ru": "ID папки для сохранения", "en": "Folder ID"},
                "Type": "string",
                "Required": "Y",
                "Multiple": "N",
            },
            "filename": {
                "Name": {"ru": "Имя файла", "en": "Filename"},
                "Type": "string",
                "Required": "Y",
                "Multiple": "N",
            },
            "variables": {
                "Name": {"ru": "Переменные шаблона (KEY|VALUE)", "en": "Template variables (KEY|VALUE)"},
                "Type": "string",
                "Required": "N",
                "Multiple": "Y",
            },
            "signatures": {
                "Name": {"ru": "Подписи (PLACEHOLDER|FILE_ID)", "en": "Signatures (PLACEHOLDER|FILE_ID)"},
                "Type": "string",
                "Required": "N",
                "Multiple": "Y",
            },
            "source_file_id": {
                "Name": {"ru": "ID существующего файла (для добавления подписи)", "en": "Source file ID (for adding signature)"},
                "Type": "string",
                "Required": "N",
                "Multiple": "N",
            },
        },
        "RETURN_PROPERTIES": {
            "file_id": {
                "Name": {"ru": "ID сгенерированного файла", "en": "Generated file ID"},
                "Type": "string",
                "Multiple": "N",
                "Default": None,
            }
        },
    }

    resp = requests.post(f"{client_endpoint}bizproc.activity.add.json", json=params)
    resp.raise_for_status()
    logger.info(f"Активити зарегистрировано: {resp.json().get('result')}")

    # Отвязываем старые пункты меню (если были привязаны с другим заголовком/хендлером)
    for old_handler in [f"{APP_URL}/form", f"{APP_URL}/app"]:
        requests.post(f"{client_endpoint}placement.unbind.json", json={
            "auth": access_token,
            "PLACEMENT": "LEFT_MENU",
            "HANDLER": old_handler,
        })

    # Регистрируем пункт в левом меню глобально
    placement_resp = requests.post(f"{client_endpoint}placement.bind.json", json={
        "auth": access_token,
        "PLACEMENT": "LEFT_MENU",
        "HANDLER": f"{APP_URL}/app",
        "TITLE": APP_TITLE,
    })
    logger.info(f"Регистрация в левом меню: {placement_resp.json()}")

    # Добавляем пункт меню каждому активному пользователю
    _bind_menu_to_all_users(access_token, client_endpoint)


async def resolve_user_id(query: dict, form_data: dict) -> tuple[str, str]:
    """Получает user_id текущего пользователя через его AUTH_ID. Возвращает (user_id, auth_id)"""
    user_id = ""
    auth_id = form_data.get("AUTH_ID", "")
    domain = query.get("DOMAIN", "") or form_data.get("DOMAIN", "")
    if auth_id and domain:
        try:
            profile_resp = requests.post(
                f"https://{domain}/rest/profile.json",
                params={"auth": auth_id}
            )
            profile = profile_resp.json().get("result", {})
            user_id = str(profile.get("ID", ""))
            logger.info(f"User ID из профиля: {user_id}")
        except Exception as e:
            logger.error(f"Ошибка получения профиля: {e}")
    return user_id, auth_id


def render_form(form_file: str, user_id: str, auth_id: str) -> str:
    """Загружает HTML формы и подставляет user_id/auth_id"""
    with open(f"/app/{form_file}", "r", encoding="utf-8") as f:
        html = f.read()
    html = html.replace("const userId = urlParams.get('user_id') || '';",
                        f"const userId = urlParams.get('user_id') || '{user_id}';")
    html = html.replace("const authToken = urlParams.get('auth_token') || '';",
                        f"const authToken = urlParams.get('auth_token') || '{auth_id}';")
    return html


def render_menu(user_id: str, auth_id: str) -> str:
    """Загружает HTML меню процессов и подставляет список процессов"""
    with open("/app/menu.html", "r", encoding="utf-8") as f:
        html = f.read()
    processes_json = json.dumps(
        {key: {"title": p["title"], "description": p["description"], "icon": p["icon"]}
         for key, p in PROCESSES.items()},
        ensure_ascii=False
    )
    html = html.replace("__APP_TITLE__", APP_TITLE)
    html = html.replace("__PROCESSES_JSON__", processes_json)
    html = html.replace("const userId = urlParams.get('user_id') || '';",
                        f"const userId = urlParams.get('user_id') || '{user_id}';")
    html = html.replace("const authToken = urlParams.get('auth_token') || '';",
                        f"const authToken = urlParams.get('auth_token') || '{auth_id}';")
    return html


# === Эндпоинты ===

@app.api_route("/install", methods=["GET", "POST"])
async def install(request: Request):
    """
    Вызывается Битрикс24 при установке приложения.
    Сохраняет токены и регистрирует активити и пункт меню.
    """
    query_params = dict(request.query_params)
    try:
        form_params = dict(await request.form())
    except Exception:
        form_params = {}

    params = {**query_params, **form_params}
    logger.info(f"Установка приложения: {params}")

    domain = params.get("DOMAIN") or params.get("auth[domain]", "")
    client_endpoint = (
        params.get("auth[client_endpoint]")
        or (f"https://{domain}/rest/" if domain else None)
    )
    tokens = {
        "access_token": params.get("AUTH_ID") or params.get("auth[access_token]"),
        "refresh_token": params.get("AUTH_REFRESH_ID") or params.get("auth[refresh_token]"),
        "domain": domain,
        "member_id": params.get("member_id") or params.get("auth[member_id]"),
        "client_endpoint": client_endpoint,
    }
    save_tokens(tokens)

    placement = params.get("PLACEMENT", "DEFAULT")
    placement_options = params.get("PLACEMENT_OPTIONS", "")

    # Страница настроек разработчика — URI содержит /devops/edit/application/
    is_devops_page = "devops" in placement_options and "edit" in placement_options

    # При открытии через LEFT_MENU или мобильное приложение — показываем меню процессов
    if placement != "DEFAULT" or not is_devops_page:
        logger.info(f"Открытие через placement={placement}, показываем меню")
        user_id, auth_id = await resolve_user_id(query_params, params)
        process_key = query_params.get("process", "")
        if process_key:
            process = PROCESSES.get(process_key)
            if process:
                return HTMLResponse(render_form(process["form_file"], user_id, auth_id))
        return HTMLResponse(render_menu(user_id, auth_id))

    # Страница настроек разработчика — регистрируем активити и показываем экран установки
    try:
        register_activity(tokens["access_token"])
    except Exception as e:
        logger.error(f"Ошибка регистрации: {e}")

    return HTMLResponse(f"""
<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>body{{font-family:sans-serif;display:flex;align-items:center;justify-content:center;height:100vh;margin:0;background:#f5f5f5;}}
.card{{background:#fff;border-radius:8px;padding:32px 40px;box-shadow:0 2px 8px rgba(0,0,0,.1);text-align:center;}}
h2{{color:#1a1a1a;margin-bottom:8px;}} p{{color:#666;font-size:14px;}}
</style></head><body>
<div class="card"><h2>✅ Приложение установлено</h2>
<p>Откройте «{APP_TITLE}» в левом меню Битрикс24.</p></div>
</body></html>""")


@app.api_route("/app", methods=["GET", "POST"], response_class=HTMLResponse)
async def app_handler(request: Request):
    """
    Обработчик приложения от Битрикс24.
    Без параметра process — показывает меню доступных процессов.
    С параметром process=<key> — открывает соответствующую форму.
    """
    query = dict(request.query_params)
    try:
        form_data = dict(await request.form())
    except Exception:
        form_data = {}

    user_id, auth_id = await resolve_user_id(query, form_data)

    process_key = query.get("process", "")

    if process_key:
        process = PROCESSES.get(process_key)
        if not process:
            return HTMLResponse(f"<p>Неизвестный процесс: {process_key}</p>", status_code=404)
        return render_form(process["form_file"], user_id, auth_id)

    return render_menu(user_id, auth_id)


@app.post("/activity-handler")
async def activity_handler(request: Request):
    """
    Обработчик активити — вызывается Битрикс24 когда БП
    доходит до действия 'Генерация документа'.
    """
    body = await request.body()
    logger.info(f"Вызов активити: {body[:200]}")

    try:
        form = await request.form()
        data = dict(form)
    except Exception:
        data = {}

    if not data:
        try:
            data = json.loads(body)
        except Exception:
            data = {}

    # Извлекаем props из form-encoded формата properties[key] и properties[key][N]
    props = {}
    for key, value in data.items():
        if not key.startswith("properties["):
            continue
        rest = key[11:]
        if "][" in rest:
            base_key = rest[:rest.index("][")]
            if base_key not in props:
                props[base_key] = []
            if isinstance(props[base_key], list):
                props[base_key].append(value)
        elif rest.endswith("]"):
            props[rest[:-1]] = value

    event_token = data.get("event_token")
    template_id = props.get("template_id", "")
    folder_id = props.get("folder_id", "")
    filename = props.get("filename", "document.docx")
    source_file_id = props.get("source_file_id", "")
    doc_data = parse_variables(props.get("variables", []))
    signatures = parse_signatures(props.get("signatures", []))

    logger.info(f"Переменные: {doc_data}")
    logger.info(f"Подписи: {signatures}")

    # Обновляем токен из данных запроса — он всегда свежий
    request_access_token = data.get("auth[access_token]")
    if request_access_token:
        tokens = load_tokens()
        tokens["access_token"] = request_access_token
        tokens["refresh_token"] = data.get("auth[refresh_token]", tokens.get("refresh_token", ""))
        tokens["client_endpoint"] = data.get("auth[client_endpoint]", tokens.get("client_endpoint", ""))
        save_tokens(tokens)

    # Форматируем REQUEST_GOAL из сырых данных "Назначение|Сумма\nНазначение2|Сумма2"
    if "REQUEST_GOAL" in doc_data:
        raw_goal = doc_data["REQUEST_GOAL"]
        lines = [l.strip() for l in raw_goal.split("\n") if l.strip()]
        if len(lines) == 1:
            # Одна позиция — берём только название без суммы
            doc_data["REQUEST_GOAL"] = lines[0].split("|")[0].strip()
        else:
            # Несколько позиций — нумерованный список
            formatted = []
            for i, line in enumerate(lines, 1):
                parts = line.split("|")
                name = parts[0].strip()
                amount = parts[1].strip() if len(parts) > 1 else ""
                if amount:
                    formatted.append(f"{i}. {name} — {int(float(amount)):,} руб.".replace(",", " "))
                else:
                    formatted.append(f"{i}. {name}")
            doc_data["REQUEST_GOAL"] = "\n" + "\n".join(formatted)

    file_id = generate_document(template_id, folder_id, filename, doc_data, signatures, source_file_id=source_file_id)

    # Возвращаем результат в БП
    client_endpoint = get_client_endpoint()
    access_token = get_access_token()
    return_url = f"{client_endpoint}bizproc.event.send.json"
    return_resp = requests.post(return_url, params={"auth": access_token}, json={
        "event_token": event_token,
        "return_values": {"file_id": file_id},
    })
    logger.info(f"Результат отправки в БП: {return_resp.text}")

    return {"status": "ok", "file_id": file_id}


@app.api_route("/form", methods=["GET", "POST"], response_class=HTMLResponse)
async def form_legacy(request: Request):
    """
    Устаревший прямой эндпоинт формы — оставлен для обратной совместимости
    со старыми ссылками. Открывает форму "Выдача наличных" напрямую.
    Новые интеграции должны использовать /app?process=<key>.
    """
    query = dict(request.query_params)
    try:
        form_data = dict(await request.form())
    except Exception:
        form_data = {}

    user_id, auth_id = await resolve_user_id(query, form_data)
    return render_form(PROCESSES["cash"]["form_file"], user_id, auth_id)


@app.post("/submit")
async def submit(req: SubmitRequest):
    """
    Принимает данные формы, создаёт элемент процесса
    и запускает бизнес-процесс в Битрикс24 от имени пользователя.
    """
    try:
        # Используем токен пользователя — тогда элемент создаётся от его имени
        access_token = req.user_auth_id if req.user_auth_id else refresh_access_token()
        client_endpoint = get_client_endpoint()

        logger.info(f"Submit: user_id={req.user_id}, title={req.title}")
        total = sum(p.amount for p in req.positions)
        positions_str = "\n".join(f"{p.name}|{int(p.amount)}" for p in req.positions)

        # Формируем красивый текст цели для документа
        if len(req.positions) == 1:
            # Одна позиция — простой текст
            request_goal = req.positions[0].name
        else:
            # Несколько позиций — нумерованный список
            lines = []
            for i, p in enumerate(req.positions, 1):
                lines.append(f"{i}. {p.name} — {int(p.amount):,} руб.".replace(",", " "))
            request_goal = "\n" + "\n".join(lines)

        # Создаём элемент процесса
        fields = {
            "NAME": req.title,
            FIELD_SUMMA: f"{int(total)}|RUB",
            FIELD_NAZNACHENIE: positions_str,
        }

        element_resp = requests.post(
            f"{client_endpoint}lists.element.add.json",
            params={"auth": access_token},
            json={
                "IBLOCK_TYPE_ID": "bitrix_processes",
                "IBLOCK_ID": CASH_REQUEST_IBLOCK_ID,
                "ELEMENT_CODE": f"cash_{datetime.now().strftime('%Y%m%d%H%M%S%f')}",
                "FIELDS": fields
            }
        )
        logger.info(f"Ответ lists.element.add: {element_resp.status_code} {element_resp.text[:300]}")
        element_resp.raise_for_status()
        element_id = element_resp.json().get("result")

        if not element_id:
            return {"success": False, "error": "Не удалось создать заявку"}

        logger.info(f"Создан элемент ID={element_id}")

        # Запускаем бизнес-процесс
        bp_resp = requests.post(
            f"{client_endpoint}bizproc.workflow.start.json",
            params={"auth": access_token},
            json={
                "TEMPLATE_ID": CASH_REQUEST_BP_ID,
                "DOCUMENT_ID": ["lists", "BizprocDocument", str(element_id)],
                "PARAMETERS": {}
            }
        )
        logger.info(f"Запуск БП: {bp_resp.json()}")

        return {"success": True, "element_id": element_id}

    except Exception as e:
        logger.error(f"Ошибка submit: {e}")
        return {"success": False, "error": str(e)}


@app.get("/health")
async def health():
    return {"status": "ok"}